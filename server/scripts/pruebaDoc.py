from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Crear un nuevo documento de Word
doc = Document()

# Añadir un título al documento
doc.add_heading('Documento con enlace a Excel', level=1)

# Añadir un párrafo con el hipervínculo al archivo de Excel
p = doc.add_paragraph()
r = p.add_run("Haz clic aquí para abrir el archivo de Excel.")

# Crear el hipervínculo
part = doc.part
r_id = part.relate_to('../Templates/SPTdt_PLTCHRNFJ.xlsx', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('r:id'), r_id)

# Añadir el hipervínculo al párrafo
new_run = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
rStyle = OxmlElement('w:rStyle')
rStyle.set(qn('w:val'), 'Hyperlink')
rPr.append(rStyle)
new_run.append(rPr)
new_run.text = " Haz clic aquí para abrir el archivo de Excel."
hyperlink.append(new_run)
p._element.append(hyperlink)

# Guardar el documento
doc.save('../documento_con_enlace.docx')
